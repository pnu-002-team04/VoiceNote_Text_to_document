import io
import os
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from google.cloud import speech
from google.cloud.speech import enums
from google.cloud.speech import types
from datetime import datetime

now = datetime.now()


def transcribe_file(speech_file):

    client = speech.SpeechClient()

    with io.open(speech_file, 'rb') as audio_file:
        content = audio_file.read()

    audio = types.RecognitionAudio(content=content)
    config = types.RecognitionConfig(
        encoding=enums.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=44100,
        language_code='en-US')

    operation = client.long_running_recognize(config, audio)

    print('Waiting for operation to complete...')
    response = operation.result(timeout=90)


    for result in response.results:
        print(u'Transcript: {}'.format(result.alternatives[0].transcript))
        f = open("./test.txt", 'w')
        f.write(result.alternatives[0].transcript)
        f.close()

        # text to doc
        document = Document()
        document.add_heading('%s-%s-%s Record File' % (now.year, now.month, now.day))
        document.add_paragraph(result.alternatives[0].transcript)
        document.save('output.docx')
        print('Confidence: {}'.format(result.alternatives[0].confidence))

if __name__ == '__main__':
    file_name = os.path.join(
        os.path.dirname(__file__),
        '.',
        'example.wav')
    with io.open(file_name, 'rb') as audio_file:
        content = audio_file.read()
        audio = types.RecognitionAudio(content=content)

    transcribe_file(file_name)


